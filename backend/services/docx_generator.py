"""ATS-safe .docx generation from resume JSON.

Rules (CLAUDE.md section 7):
- Font: Calibri only, body 11pt, headings 13pt, name 16pt
- Single column — no tables, text boxes, or frames
- Section headings: SUMMARY, EXPERIENCE, SKILLS, EDUCATION (uppercase)
- Separator between title and company: pipe character | not em dash
- No colour except pure black RGB(0,0,0)
- No images, icons, or graphics of any kind
- Margins: 0.5 inch top/bottom, 0.75 inch left/right
- Generate on demand — do not store permanently
"""
import logging
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────

_FONT = "Calibri"
_SIZE_NAME = Pt(16)
_SIZE_HEADING = Pt(13)
_SIZE_BODY = Pt(11)
_BLACK = RGBColor(0x00, 0x00, 0x00)
_MARGIN_TB = Inches(0.5)
_MARGIN_LR = Inches(0.75)


# ── Internal helpers ──────────────────────────────────────────────────────────

def _set_run_style(run, size: Pt, bold: bool = False) -> None:
    """Apply font, size, colour, and bold to a run. All formatting explicit."""
    run.font.name = _FONT
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = _BLACK
    # Suppress theme-colour override
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _FONT)
    rFonts.set(qn("w:hAnsi"), _FONT)
    rFonts.set(qn("w:cs"), _FONT)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _clear_para_spacing(para) -> None:
    """Remove before/after spacing so sections sit tight."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"), "60")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def _add_para(doc: Document, text: str, size: Pt, bold: bool = False, space_before: int = 0) -> None:
    """Add a plain paragraph with explicit Calibri formatting."""
    para = doc.add_paragraph()
    _clear_para_spacing(para)
    if space_before:
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:before"), str(space_before))
    run = para.add_run(text)
    _set_run_style(run, size, bold=bold)


def _add_section_heading(doc: Document, text: str) -> None:
    """Add an uppercase section heading (13pt bold, full-width underline via bottom border)."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    # Bottom border to create a visual rule under the heading
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    # Spacing above heading
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "160")
    spacing.set(qn("w:after"), "40")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)
    run = para.add_run(text.upper())
    _set_run_style(run, _SIZE_HEADING, bold=True)


def _add_bullet(doc: Document, text: str) -> None:
    """Add a plain bullet point using an em-dash prefix (no list XML — ATS-safe)."""
    para = doc.add_paragraph()
    _clear_para_spacing(para)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)
    run = para.add_run(f"•  {text}")
    _set_run_style(run, _SIZE_BODY)


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_ats_resume_docx(resume_json: dict) -> bytes:
    """Generate an ATS-safe .docx from resume JSON and return the raw bytes.

    The caller streams the bytes directly — no file is saved to disk.

    Args:
        resume_json: dict matching the ResumeJson schema (full_name, contact,
                     summary, experiences, skills, education).

    Returns:
        Raw .docx bytes ready to stream as a response.
    """
    doc = Document()

    # ── Margins ───────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = _MARGIN_TB
        section.bottom_margin = _MARGIN_TB
        section.left_margin = _MARGIN_LR
        section.right_margin = _MARGIN_LR

    # Remove default empty paragraph Word adds to new documents
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # ── Name ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    _clear_para_spacing(name_para)
    name_run = name_para.add_run(resume_json.get("full_name", ""))
    _set_run_style(name_run, _SIZE_NAME, bold=True)
    name_para.alignment = 1  # CENTER

    # ── Contact ───────────────────────────────────────────────────────────────
    contact = resume_json.get("contact", "").strip()
    if contact:
        contact_para = doc.add_paragraph()
        _clear_para_spacing(contact_para)
        contact_run = contact_para.add_run(contact)
        _set_run_style(contact_run, _SIZE_BODY)
        contact_para.alignment = 1  # CENTER

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    summary = resume_json.get("summary", "").strip()
    if summary:
        _add_section_heading(doc, "Summary")
        _add_para(doc, summary, _SIZE_BODY)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    experiences = resume_json.get("experiences") or []
    if experiences:
        _add_section_heading(doc, "Experience")
        for exp in experiences:
            title = exp.get("title", "")
            company = exp.get("company", "")
            dates = exp.get("dates", "")
            bullets = exp.get("bullets") or []

            # "Title | Company" on one line, dates on the same or next
            title_company = f"{title} | {company}" if title and company else title or company
            entry_para = doc.add_paragraph()
            _clear_para_spacing(entry_para)
            pPr = entry_para._p.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "120")
            spacing.set(qn("w:after"), "40")
            existing = pPr.find(qn("w:spacing"))
            if existing is not None:
                pPr.remove(existing)
            pPr.append(spacing)

            # Title | Company — bold
            tc_run = entry_para.add_run(title_company)
            _set_run_style(tc_run, _SIZE_BODY, bold=True)

            # Dates — right-aligned via tab stop is ATS-risky; use separator instead
            if dates:
                sep_run = entry_para.add_run(f"  —  {dates}")
                _set_run_style(sep_run, _SIZE_BODY, bold=False)

            for bullet in bullets:
                _add_bullet(doc, bullet)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    skills = resume_json.get("skills") or {}
    technical = skills.get("technical") or []
    tools = skills.get("tools") or []
    if technical or tools:
        _add_section_heading(doc, "Skills")
        if technical:
            _add_para(doc, "Technical:  " + ",  ".join(technical), _SIZE_BODY)
        if tools:
            _add_para(doc, "Tools & Platforms:  " + ",  ".join(tools), _SIZE_BODY)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    education = resume_json.get("education") or []
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year") or ""
            parts = [p for p in [degree, institution, year] if p]
            if parts:
                _add_para(doc, "  |  ".join(parts), _SIZE_BODY)

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()
    logger.info(
        "Generated ATS resume docx: %d bytes for '%s'",
        len(docx_bytes), resume_json.get("full_name", "unknown"),
    )
    return docx_bytes
